import os
import tempfile
import asyncio
from typing import Dict, Any, List, Optional
from docx import Document
from pypdf import PdfReader
import extract_msg
import aiofiles


async def parse_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n".join(text_parts)


async def parse_docx(file_path: str) -> str:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                tables_text.append(row_text)
    
    return "\n".join(paragraphs + tables_text)


async def parse_msg(file_path: str) -> str:
    msg = extract_msg.Message(file_path)
    parts = []
    
    if msg.subject:
        parts.append(f"Тема: {msg.subject}")
    if msg.sender:
        parts.append(f"От: {msg.sender}")
    if msg.date:
        parts.append(f"Дата: {msg.date}")
    if msg.body:
        parts.append(f"\n{msg.body}")
    
    return "\n".join(parts)


async def parse_txt(file_path: str) -> str:
    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return await f.read()


async def parse_file(file_content: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name
    
    try:
        if ext == '.pdf':
            return await parse_pdf(tmp_path)
        elif ext in ['.docx', '.doc']:
            return await parse_docx(tmp_path)
        elif ext == '.msg':
            return await parse_msg(tmp_path)
        elif ext == '.txt':
            return await parse_txt(tmp_path)
        else:
            return file_content.decode('utf-8', errors='ignore')
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def get_allowed_extensions() -> List[str]:
    return ['.pdf', '.docx', '.doc', '.msg', '.txt']
